"""
create_test_fixture.py — Generate a test .docx file with various formatting
and math equations for testing the DOCX-to-DOCX pipeline.

Run this script to generate: tests/fixtures/test_formatted.docx
"""

import os
import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


def create_test_docx(output_path: str) -> None:
    """Create a test .docx with diverse formatting and math equations."""
    doc = Document()

    # ── 1. Title paragraph (centered, large, bold) ──────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Test Document for doc-extract-reconstruct")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue

    # ── 2. Normal paragraph with mixed formatting ───────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run("This is ")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    run = p.add_run("bold text")
    run.bold = True
    run.font.size = Pt(12)

    run = p.add_run(", ")
    run.font.size = Pt(12)

    run = p.add_run("italic text")
    run.italic = True
    run.font.size = Pt(12)

    run = p.add_run(", ")
    run.font.size = Pt(12)

    run = p.add_run("underlined text")
    run.underline = True
    run.font.size = Pt(12)

    run = p.add_run(", and ")
    run.font.size = Pt(12)

    run = p.add_run("bold italic")
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)

    run = p.add_run(".")
    run.font.size = Pt(12)

    # ── 3. Colored text paragraph ────────────────────────────────
    p = doc.add_paragraph()
    colors = [
        ("Red text ", RGBColor(0xFF, 0x00, 0x00)),
        ("Green text ", RGBColor(0x00, 0x80, 0x00)),
        ("Blue text ", RGBColor(0x00, 0x00, 0xFF)),
        ("Orange text", RGBColor(0xFF, 0x80, 0x00)),
    ]
    for text, color in colors:
        run = p.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)

    # ── 4. Different fonts paragraph ─────────────────────────────
    p = doc.add_paragraph()
    fonts = [
        ("Times New Roman, ", "Times New Roman"),
        ("Arial, ", "Arial"),
        ("Courier New, ", "Courier New"),
        ("Calibri", "Calibri"),
    ]
    for text, font in fonts:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(11)

    # ── 5. Different sizes paragraph ─────────────────────────────
    p = doc.add_paragraph()
    sizes = [8, 10, 12, 14, 18, 24]
    for size in sizes:
        run = p.add_run(f"{size}pt ")
        run.font.size = Pt(size)

    # ── 6. Math equation (OMML) ──────────────────────────────────
    # Insert a native Word math equation using raw OMML XML.
    # This is the Pythagorean theorem: a² + b² = c²
    p = doc.add_paragraph()
    run = p.add_run("The Pythagorean theorem: ")
    run.font.size = Pt(12)

    # Build the OMML element for a² + b² = c²
    omml_xml = _create_pythagorean_omml()
    p._element.append(omml_xml)

    # ── 7. Another math equation (quadratic formula) ─────────────
    p = doc.add_paragraph()
    run = p.add_run("Quadratic formula: ")
    run.font.size = Pt(12)

    omml_xml = _create_quadratic_omml()
    p._element.append(omml_xml)

    # ── 8. Aligned paragraph (right) ─────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("This paragraph is right-aligned.")
    run.font.size = Pt(11)

    # ── 9. Justified paragraph ───────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "This is a justified paragraph. It should stretch across the full width "
        "of the page margins. The text alignment property should be preserved "
        "during reconstruction from the source document to the output document."
    )
    run.font.size = Pt(11)

    # ── 10. Bullet list ──────────────────────────────────────────
    doc.add_paragraph("First bullet item", style='List Bullet')
    doc.add_paragraph("Second bullet item", style='List Bullet')
    doc.add_paragraph("Third bullet item", style='List Bullet')

    # ── 11. Numbered list ────────────────────────────────────────
    doc.add_paragraph("First numbered item", style='List Number')
    doc.add_paragraph("Second numbered item", style='List Number')
    doc.add_paragraph("Third numbered item", style='List Number')

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Test fixture created: {output_path}")


def _create_pythagorean_omml() -> etree._Element:
    """
    Create OMML XML for: a² + b² = c²
    Uses the Office MathML namespace.
    """
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    nsmap = {'m': M}

    oMath = etree.Element(f'{{{M}}}oMath', nsmap=nsmap)

    # a²
    _add_math_run(oMath, M, "a")
    _add_superscript(oMath, M, "2")

    # +
    _add_math_run(oMath, M, "+")

    # b²
    _add_math_run(oMath, M, "b")
    _add_superscript(oMath, M, "2")

    # =
    _add_math_run(oMath, M, "=")

    # c²
    _add_math_run(oMath, M, "c")
    _add_superscript(oMath, M, "2")

    return oMath


def _create_quadratic_omml() -> etree._Element:
    """
    Create OMML XML for the quadratic formula:
    x = (-b ± √(b² - 4ac)) / 2a

    This is a more complex equation to test fraction and radical handling.
    """
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    nsmap = {'m': M, 'w': W}

    oMath = etree.Element(f'{{{M}}}oMath', nsmap=nsmap)

    # x =
    _add_math_run(oMath, M, "x=")

    # Fraction: (-b ± √(b²-4ac)) / (2a)
    f = etree.SubElement(oMath, f'{{{M}}}f')

    # Fraction properties
    fPr = etree.SubElement(f, f'{{{M}}}fPr')

    # Numerator: -b ± √(b²-4ac)
    num = etree.SubElement(f, f'{{{M}}}num')
    _add_math_run(num, M, "-b±√(b²-4ac)")

    # Denominator: 2a
    den = etree.SubElement(f, f'{{{M}}}den')
    _add_math_run(den, M, "2a")

    return oMath


def _add_math_run(parent: etree._Element, ns: str, text: str) -> None:
    """Add a math run (m:r) with text to a parent element."""
    r = etree.SubElement(parent, f'{{{ns}}}r')
    t = etree.SubElement(r, f'{{{ns}}}t')
    t.text = text


def _add_superscript(parent: etree._Element, ns: str, text: str) -> None:
    """Add a superscript element to a parent."""
    sSup = etree.SubElement(parent, f'{{{ns}}}sSup')

    # Base (empty — the previous run serves as the visual base)
    e = etree.SubElement(sSup, f'{{{ns}}}e')

    # Superscript content
    sup = etree.SubElement(sSup, f'{{{ns}}}sup')
    _add_math_run(sup, ns, text)


# ── Run ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    output = os.path.join(
        os.path.dirname(__file__),
        "fixtures",
        "test_formatted.docx",
    )
    create_test_fixture = create_test_docx
    create_test_fixture(output)
