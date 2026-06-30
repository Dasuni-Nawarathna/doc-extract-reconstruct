"""
test_docx_pipeline.py — Tests for the DOCX-to-DOCX reconstruction pipeline.

Verifies that text formatting (font, size, bold, italic, underline, color),
paragraph properties (alignment, spacing), and math equations (OMML) are
preserved through the extraction-reconstruction cycle.
"""

import os
import sys
import pytest
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from lxml import etree

from doc_extract_reconstruct.docx_pipeline import process_docx
from doc_extract_reconstruct.writer import qn
from doc_extract_reconstruct.router import detect_input_type, InputType


# ── Fixtures ───────────────────────────────────────────────────────────

FIXTURES_DIR = Path(__file__).parent / "fixtures"
INPUT_DOCX = FIXTURES_DIR / "test_formatted.docx"
OUTPUT_DOCX = FIXTURES_DIR / "test_output.docx"


@pytest.fixture(autouse=True)
def cleanup():
    """Clean up output files after each test."""
    yield
    for f in FIXTURES_DIR.glob("test_output*"):
        f.unlink(missing_ok=True)


def _get_all_runs(doc: Document) -> list:
    """Extract all runs from a document with their formatting."""
    runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            })
    return runs


def _count_omml_elements(doc: Document) -> int:
    """Count the number of OMML elements in a document."""
    body = doc.element.body
    m_omath = qn('m:oMath')
    m_omathpara = qn('m:oMathPara')
    count = 0
    for elem in body.iter():
        if elem.tag in (m_omath, m_omathpara):
            # Don't count nested elements
            parent = elem.getparent()
            if parent is not None and parent.tag in (m_omath, m_omathpara):
                continue
            count += 1
    return count


# ── Tests ──────────────────────────────────────────────────────────────

class TestRouter:
    """Test input type detection."""

    def test_detect_docx(self):
        assert detect_input_type(str(INPUT_DOCX)) == InputType.DOCX

    def test_detect_unsupported(self):
        with pytest.raises(ValueError, match="Unsupported file extension"):
            detect_input_type("test.xyz")

    def test_detect_missing_file(self):
        with pytest.raises(FileNotFoundError):
            detect_input_type("nonexistent.docx")


class TestDocxPipeline:
    """Test the DOCX-to-DOCX pipeline."""

    @pytest.fixture(autouse=True)
    def run_pipeline(self):
        """Run the pipeline once for all tests in this class."""
        assert INPUT_DOCX.exists(), f"Test fixture not found: {INPUT_DOCX}"
        process_docx(str(INPUT_DOCX), str(OUTPUT_DOCX))
        self.src_doc = Document(str(INPUT_DOCX))
        self.dst_doc = Document(str(OUTPUT_DOCX))
        yield
        # cleanup handled by the module-level fixture

    def test_output_file_created(self):
        """Verify the output file was created."""
        assert OUTPUT_DOCX.exists()

    def test_paragraph_count_preserved(self):
        """Verify the same number of paragraphs in source and output."""
        src_count = len(self.src_doc.paragraphs)
        dst_count = len(self.dst_doc.paragraphs)
        assert dst_count == src_count, (
            f"Paragraph count mismatch: source={src_count}, output={dst_count}"
        )

    def test_text_content_preserved(self):
        """Verify all text content is preserved."""
        src_text = "\n".join(p.text for p in self.src_doc.paragraphs)
        dst_text = "\n".join(p.text for p in self.dst_doc.paragraphs)
        assert dst_text == src_text, "Text content does not match"

    def test_bold_preserved(self):
        """Verify bold formatting is preserved."""
        src_runs = _get_all_runs(self.src_doc)
        dst_runs = _get_all_runs(self.dst_doc)

        # Find bold runs in source
        src_bold = [(r['text'], r['bold']) for r in src_runs if r['bold']]
        dst_bold = [(r['text'], r['bold']) for r in dst_runs if r['bold']]

        assert len(dst_bold) >= len(src_bold), (
            f"Bold runs: source has {len(src_bold)}, output has {len(dst_bold)}"
        )

    def test_italic_preserved(self):
        """Verify italic formatting is preserved."""
        src_runs = _get_all_runs(self.src_doc)
        dst_runs = _get_all_runs(self.dst_doc)

        src_italic = [r['text'] for r in src_runs if r['italic']]
        dst_italic = [r['text'] for r in dst_runs if r['italic']]

        assert len(dst_italic) >= len(src_italic), (
            f"Italic runs: source has {len(src_italic)}, output has {len(dst_italic)}"
        )

    def test_font_colors_preserved(self):
        """Verify font colors are preserved."""
        src_runs = _get_all_runs(self.src_doc)
        dst_runs = _get_all_runs(self.dst_doc)

        src_colored = [(r['text'], str(r['font_color'])) for r in src_runs if r['font_color']]
        dst_colored = [(r['text'], str(r['font_color'])) for r in dst_runs if r['font_color']]

        assert len(dst_colored) >= len(src_colored), (
            f"Colored runs: source has {len(src_colored)}, output has {len(dst_colored)}"
        )

    def test_font_sizes_preserved(self):
        """Verify font sizes are preserved."""
        src_runs = _get_all_runs(self.src_doc)
        dst_runs = _get_all_runs(self.dst_doc)

        src_sizes = {r['text']: r['font_size'] for r in src_runs if r['font_size']}
        dst_sizes = {r['text']: r['font_size'] for r in dst_runs if r['font_size']}

        for text, size in src_sizes.items():
            if text in dst_sizes:
                assert dst_sizes[text] == size, (
                    f"Font size mismatch for '{text}': "
                    f"source={size}, output={dst_sizes[text]}"
                )

    def test_omml_equations_preserved(self):
        """Verify OMML math equations are preserved."""
        src_count = _count_omml_elements(self.src_doc)
        dst_count = _count_omml_elements(self.dst_doc)

        assert src_count > 0, "Source document should contain OMML equations"
        assert dst_count == src_count, (
            f"OMML equation count mismatch: source={src_count}, output={dst_count}"
        )

    def test_alignment_preserved(self):
        """Verify paragraph alignment is preserved in the XML."""
        src_body = self.src_doc.element.body
        dst_body = self.dst_doc.element.body

        src_alignments = []
        dst_alignments = []

        for p in src_body.findall(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    src_alignments.append(jc.get(qn('w:val')))

        for p in dst_body.findall(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    dst_alignments.append(jc.get(qn('w:val')))

        assert dst_alignments == src_alignments, (
            f"Alignment mismatch: source={src_alignments}, output={dst_alignments}"
        )

    def test_confidence_report_created(self):
        """Verify a confidence report was generated."""
        report_path = FIXTURES_DIR / "test_output_confidence_report.txt"
        assert report_path.exists(), "Confidence report was not created"

        content = report_path.read_text()
        assert "DOC-EXTRACT-RECONSTRUCT" in content
        assert "HIGH confidence" in content


class TestDocxPipelineXml:
    """Lower-level XML tests for the DOCX pipeline."""

    @pytest.fixture(autouse=True)
    def run_pipeline(self):
        assert INPUT_DOCX.exists()
        process_docx(str(INPUT_DOCX), str(OUTPUT_DOCX))
        self.src_doc = Document(str(INPUT_DOCX))
        self.dst_doc = Document(str(OUTPUT_DOCX))
        yield

    def test_rfonts_preserved(self):
        """Verify w:rFonts elements are preserved in runs."""
        src_body = self.src_doc.element.body
        dst_body = self.dst_doc.element.body

        src_fonts = set()
        for rFonts in src_body.iter(qn('w:rFonts')):
            ascii_font = rFonts.get(qn('w:ascii'))
            if ascii_font:
                src_fonts.add(ascii_font)

        dst_fonts = set()
        for rFonts in dst_body.iter(qn('w:rFonts')):
            ascii_font = rFonts.get(qn('w:ascii'))
            if ascii_font:
                dst_fonts.add(ascii_font)

        assert src_fonts.issubset(dst_fonts), (
            f"Missing fonts in output: {src_fonts - dst_fonts}"
        )

    def test_omml_xml_structure_preserved(self):
        """Verify the OMML XML structure is identical in source and output."""
        m_omath = qn('m:oMath')

        src_maths = list(self.src_doc.element.body.iter(m_omath))
        dst_maths = list(self.dst_doc.element.body.iter(m_omath))

        # Compare XML serialization of each math element
        for src_m, dst_m in zip(src_maths, dst_maths):
            src_xml = etree.tostring(src_m, pretty_print=True).decode()
            dst_xml = etree.tostring(dst_m, pretty_print=True).decode()
            assert src_xml == dst_xml, (
                f"OMML XML mismatch:\nSource:\n{src_xml}\nOutput:\n{dst_xml}"
            )


# ── Run ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
